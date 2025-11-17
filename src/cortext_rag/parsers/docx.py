"""Word document parser."""

from pathlib import Path
from typing import Any


class DOCXParser:
    """Parse DOCX files using python-docx."""

    @property
    def supported_extensions(self) -> list[str]:
        """Supported file extensions."""
        return [".docx"]

    def parse(self, path: Path) -> tuple[str, dict[str, Any]]:
        """Parse DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            Tuple of (content, metadata)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not installed. "
                "Install with: pip install cortext-workspace[rag]"
            )

        doc = Document(path)

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        # Extract metadata
        metadata = {
            "num_paragraphs": len(doc.paragraphs),
            "file_name": path.name,
        }

        # Add core properties if available
        if doc.core_properties:
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author

        return content.strip(), metadata
